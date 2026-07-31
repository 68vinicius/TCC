"""

streamlit run src/app/frontend.py

"""

from __future__ import annotations

import sys
from pathlib import Path

_ROOT_DIR = Path(__file__).resolve().parent.parent
if str(_ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(_ROOT_DIR))

import io
import csv
import time
import string
import hashlib
import logging
import itertools
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
import streamlit as st

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# CONFIGURAÇÃO DA PÁGINA 

st.set_page_config(
    page_title="Sistema Computacional de Auditoria Textual",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        "Get Help": "https://github.com/68vinicius/TCC",
        "About": "Sistema Computacional de Auditoria Textual",
    },
)

GLOBAL_CSS = """
<style>
/* ── Google Fonts ─────────────────────────────────── */
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600;700&family=Inter:wght@300;400;500;600&display=swap');

/* ── Tokens ────────────────────────────────────────── */
:root {
  --bg-base:       #0D1117;
  --bg-surface:    #161B22;
  --bg-elevated:   #1C2128;
  --bg-card:       #21262D;
  --border:        #30363D;
  --border-accent: #00D4FF40;
  --text-primary:  #E6EDF3;
  --text-muted:    #8B949E;
  --text-faint:    #484F58;
  --accent:        #00D4FF;
  --accent-dim:    #00D4FF20;
  --accent-mid:    #0EA5E9;
  --purple:        #7C3AED;
  --purple-dim:    #7C3AED20;
  --green:         #3FB950;
  --green-dim:     #3FB95020;
  --yellow:        #D29922;
  --red:           #F85149;
  --red-dim:       #F8514920;
  --radius:        8px;
  --radius-lg:     12px;
  --font-mono:     'IBM Plex Mono', 'Fira Code', monospace;
  --font-body:     'Inter', -apple-system, sans-serif;
}

/* ── Reset / Base ──────────────────────────────────── */
html, body, [class*="css"] {
  background-color: var(--bg-base) !important;
  color: var(--text-primary) !important;
  font-family: var(--font-body) !important;
}

/* ── Oculta elementos default do Streamlit ─────────── */
#MainMenu, footer, header { visibility: hidden; }
.stDeployButton { display: none; }

/* ── Scrollbar ─────────────────────────────────────── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--bg-base); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: var(--accent); }

/* ── Layout principal ──────────────────────────────── */
.main .block-container {
  padding: 2rem 2.5rem 4rem !important;
  max-width: 1400px !important;
}

/* ── Sidebar ───────────────────────────────────────── */
[data-testid="stSidebar"] {
  background-color: var(--bg-surface) !important;
  border-right: 1px solid var(--border) !important;
}
[data-testid="stSidebar"] .stMarkdown p,
[data-testid="stSidebar"] label {
  color: var(--text-muted) !important;
  font-size: 0.82rem !important;
}

/* ── Header hero ───────────────────────────────────── */
.hero-container {
  background: linear-gradient(135deg, var(--bg-surface) 0%, var(--bg-elevated) 100%);
  border: 1px solid var(--border);
  border-top: 3px solid var(--accent);
  border-radius: var(--radius-lg);
  padding: 2.2rem 2.5rem 1.8rem;
  margin-bottom: 2rem;
  position: relative;
  overflow: hidden;
}
.hero-container::before {
  content: '';
  position: absolute;
  top: 0; right: 0;
  width: 300px; height: 300px;
  background: radial-gradient(circle at 100% 0%, var(--accent-dim) 0%, transparent 70%);
  pointer-events: none;
}
.hero-title {
  font-family: var(--font-mono) !important;
  font-size: 1.85rem !important;
  font-weight: 700 !important;
  color: var(--text-primary) !important;
  letter-spacing: -0.03em;
  margin: 0 0 0.4rem 0;
}
.hero-title span { color: var(--accent); }
.hero-subtitle {
  color: var(--text-muted) !important;
  font-size: 0.88rem !important;
  font-weight: 400 !important;
  letter-spacing: 0.02em;
}
.hero-badge {
  display: inline-block;
  background: var(--accent-dim);
  border: 1px solid var(--accent);
  color: var(--accent);
  font-family: var(--font-mono);
  font-size: 0.70rem;
  padding: 0.15rem 0.6rem;
  border-radius: 3px;
  margin-top: 0.8rem;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

/* ── Section headers ───────────────────────────────── */
.section-label {
  font-family: var(--font-mono);
  font-size: 0.70rem;
  font-weight: 600;
  color: var(--accent);
  letter-spacing: 0.12em;
  text-transform: uppercase;
  margin-bottom: 0.9rem;
  display: flex;
  align-items: center;
  gap: 0.5rem;
}
.section-label::after {
  content: '';
  flex: 1;
  height: 1px;
  background: var(--border);
}

/* ── Upload zone ───────────────────────────────────── */
[data-testid="stFileUploader"] {
  background: var(--bg-surface) !important;
  border: 1.5px dashed var(--border) !important;
  border-radius: var(--radius-lg) !important;
  padding: 1rem !important;
  transition: border-color 0.2s;
}
[data-testid="stFileUploader"]:hover {
  border-color: var(--accent) !important;
}
[data-testid="stFileUploader"] label {
  color: var(--text-primary) !important;
  font-size: 0.88rem !important;
}

/* ── File cards ─────────────────────────────────────── */
.file-card {
  background: var(--bg-card);
  border: 1px solid var(--border);
  border-left: 3px solid var(--accent);
  border-radius: var(--radius);
  padding: 0.75rem 1rem;
  margin-bottom: 0.5rem;
  display: flex;
  align-items: center;
  gap: 0.75rem;
  transition: border-color 0.15s, background 0.15s;
}
.file-card:hover { border-color: var(--accent); background: var(--bg-elevated); }
.file-icon { font-size: 1.2rem; }
.file-info { flex: 1; }
.file-name {
  font-family: var(--font-mono);
  font-size: 0.82rem;
  font-weight: 600;
  color: var(--text-primary);
}
.file-meta {
  font-size: 0.72rem;
  color: var(--text-muted);
  margin-top: 0.1rem;
}
.file-badge {
  font-family: var(--font-mono);
  font-size: 0.65rem;
  padding: 0.1rem 0.4rem;
  border-radius: 3px;
  text-transform: uppercase;
  font-weight: 600;
}
.badge-pdf  { background: #F8514920; color: #F85149; border: 1px solid #F8514940; }
.badge-docx { background: #0EA5E920; color: #0EA5E9; border: 1px solid #0EA5E940; }
.badge-txt  { background: #3FB95020; color: #3FB950; border: 1px solid #3FB95040; }

/* ── Metric cards ───────────────────────────────────── */
.metric-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
  gap: 0.75rem;
  margin-bottom: 1.5rem;
}
.metric-card {
  background: var(--bg-card);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 1rem 1.1rem;
  position: relative;
  overflow: hidden;
}
.metric-card::before {
  content: '';
  position: absolute;
  bottom: 0; left: 0; right: 0;
  height: 2px;
}
.metric-card.cyan::before   { background: var(--accent); }
.metric-card.purple::before { background: var(--purple); }
.metric-card.green::before  { background: var(--green); }
.metric-card.yellow::before { background: var(--yellow); }
.metric-label {
  font-family: var(--font-mono);
  font-size: 0.65rem;
  color: var(--text-muted);
  letter-spacing: 0.08em;
  text-transform: uppercase;
  margin-bottom: 0.35rem;
}
.metric-value {
  font-family: var(--font-mono);
  font-size: 1.6rem;
  font-weight: 700;
  color: var(--text-primary);
  line-height: 1;
}
.metric-sub {
  font-size: 0.72rem;
  color: var(--text-faint);
  margin-top: 0.25rem;
}

/* ── Risk pill ──────────────────────────────────────── */
.risk-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.4rem;
  padding: 0.25rem 0.75rem;
  border-radius: 999px;
  font-family: var(--font-mono);
  font-size: 0.72rem;
  font-weight: 600;
  letter-spacing: 0.05em;
}
.risk-high    { background: var(--red-dim);    color: var(--red);    border: 1px solid #F8514960; }
.risk-elevado { background: #D2992220;         color: var(--yellow); border: 1px solid #D2992260; }
.risk-mod     { background: var(--accent-dim); color: var(--accent); border: 1px solid #00D4FF60; }
.risk-low     { background: var(--green-dim);  color: var(--green);  border: 1px solid #3FB95060; }

/* ── Result table ───────────────────────────────────── */
.result-row {
  background: var(--bg-card);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 0.9rem 1.1rem;
  margin-bottom: 0.6rem;
  display: grid;
  grid-template-columns: auto 1fr auto auto;
  align-items: center;
  gap: 1rem;
  transition: background 0.15s;
}
.result-row:hover { background: var(--bg-elevated); }
.rank-badge {
  font-family: var(--font-mono);
  font-size: 0.75rem;
  font-weight: 700;
  color: var(--text-faint);
  min-width: 1.8rem;
  text-align: center;
}
.rank-1 { color: var(--accent); }
.rank-2 { color: var(--purple); }
.rank-3 { color: var(--green); }
.pair-names {
  font-family: var(--font-mono);
  font-size: 0.80rem;
  color: var(--text-primary);
}
.pair-names span { color: var(--text-muted); }

/* ── Score bar ──────────────────────────────────────── */
.score-bar-wrap { min-width: 140px; }
.score-bar-track {
  background: var(--bg-base);
  border-radius: 999px;
  height: 6px;
  overflow: hidden;
  margin-bottom: 0.2rem;
}
.score-bar-fill {
  height: 100%;
  border-radius: 999px;
  transition: width 0.6s ease;
}
.score-pct {
  font-family: var(--font-mono);
  font-size: 0.78rem;
  font-weight: 600;
  color: var(--text-primary);
  text-align: right;
}

/* ── Snippet box ────────────────────────────────────── */
.snippet-box {
  background: var(--bg-base);
  border: 1px solid var(--border);
  border-radius: var(--radius);
  padding: 0.9rem 1rem;
  font-family: var(--font-mono);
  font-size: 0.78rem;
  color: var(--text-muted);
  line-height: 1.7;
  white-space: pre-wrap;
  word-break: break-word;
  max-height: 200px;
  overflow-y: auto;
}
.snippet-box mark {
  background: var(--accent-dim);
  color: var(--accent);
  padding: 0.05em 0.2em;
  border-radius: 2px;
}

/* ── Expander ───────────────────────────────────────── */
[data-testid="stExpander"] {
  background: var(--bg-card) !important;
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  margin-bottom: 0.5rem;
}
[data-testid="stExpander"] summary {
  font-family: var(--font-mono) !important;
  font-size: 0.82rem !important;
  color: var(--text-primary) !important;
  padding: 0.75rem 1rem !important;
}

/* ── Buttons ────────────────────────────────────────── */
.stButton > button {
  background: var(--accent) !important;
  color: #000 !important;
  font-family: var(--font-mono) !important;
  font-size: 0.82rem !important;
  font-weight: 700 !important;
  letter-spacing: 0.06em !important;
  border: none !important;
  border-radius: var(--radius) !important;
  padding: 0.65rem 1.5rem !important;
  cursor: pointer !important;
  transition: opacity 0.15s, transform 0.1s !important;
}
.stButton > button:hover {
  opacity: 0.88 !important;
  transform: translateY(-1px) !important;
}
.stButton > button:active { transform: translateY(0) !important; }

/* ── Download button ────────────────────────────────── */
.stDownloadButton > button {
  background: var(--bg-card) !important;
  color: var(--text-primary) !important;
  border: 1px solid var(--border) !important;
  font-family: var(--font-mono) !important;
  font-size: 0.78rem !important;
}
.stDownloadButton > button:hover {
  border-color: var(--accent) !important;
  color: var(--accent) !important;
}

/* ── Progress bar ───────────────────────────────────── */
[data-testid="stProgress"] > div > div {
  background: linear-gradient(90deg, var(--accent), var(--purple)) !important;
  border-radius: 999px !important;
}
[data-testid="stProgress"] > div {
  background: var(--bg-card) !important;
  border-radius: 999px !important;
  height: 6px !important;
}

/* ── Dataframe ──────────────────────────────────────── */
[data-testid="stDataFrame"] {
  border: 1px solid var(--border) !important;
  border-radius: var(--radius) !important;
  overflow: hidden;
}
[data-testid="stDataFrame"] table {
  font-family: var(--font-mono) !important;
  font-size: 0.78rem !important;
}

/* ── Divider ────────────────────────────────────────── */
hr {
  border: none !important;
  border-top: 1px solid var(--border) !important;
  margin: 1.5rem 0 !important;
}

/* ── Alerts / Callouts ──────────────────────────────── */
[data-testid="stAlert"] {
  background: var(--bg-card) !important;
  border-radius: var(--radius) !important;
  font-family: var(--font-mono) !important;
  font-size: 0.80rem !important;
}

/* ── Checkbox / Radio ────────────────────────────────── */
.stCheckbox label, .stRadio label {
  font-size: 0.82rem !important;
  color: var(--text-muted) !important;
}

/* ── Selectbox ──────────────────────────────────────── */
[data-testid="stSelectbox"] select,
[data-testid="stSelectbox"] > div > div {
  background: var(--bg-card) !important;
  border: 1px solid var(--border) !important;
  color: var(--text-primary) !important;
  font-family: var(--font-mono) !important;
  font-size: 0.82rem !important;
  border-radius: var(--radius) !important;
}

/* ── Tooltip ────────────────────────────────────────── */
.tooltip-box {
  background: var(--bg-elevated);
  border: 1px solid var(--border-accent);
  border-radius: var(--radius);
  padding: 0.6rem 0.9rem;
  font-family: var(--font-mono);
  font-size: 0.75rem;
  color: var(--text-muted);
  margin-top: 0.5rem;
}
</style>
"""

# CAMADA DE EXTRAÇÃO DE DOCUMENTOS

def _try_import_factory():
    """
    Tenta importar DocumentFactory do backend do TCC.
    Retorna a classe se disponível, None caso contrário.
    """
    try:
        import sys
        from pathlib import Path
        root = Path(__file__).resolve().parent
        if str(root) not in sys.path:
            sys.path.insert(0, str(root))
        from src.processamento.extrator import DocumentFactory
        return DocumentFactory
    except ImportError:
        return None


def extract_text(uploaded_file) -> str:
    """
    Extrai texto de um arquivo carregado via Streamlit.

    Prioridade:
      1. DocumentFactory (src/processamento/extrator.py) — backend completo do TCC,
         com suporte a metadados, detecção de encoding robusta e logging.
      2. Fallback embutido — garante funcionamento mesmo sem o backend instalado.

    Esta função é o único ponto de entrada de extração no frontend,
    eliminando duplicação de lógica e garantindo consistência.
    """
    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()

    DocumentFactory = _try_import_factory()
    if DocumentFactory is not None:
        try:
            uploaded_file.seek(0)
            doc = DocumentFactory.extract(uploaded_file, uploaded_file.name)
            uploaded_file.seek(0)
            return doc.raw_text
        except Exception as e:
            logger.warning("DocumentFactory falhou (%s), usando fallback: %s", uploaded_file.name, e)

    uploaded_file.seek(0)
    raw = uploaded_file.read()
    uploaded_file.seek(0)

    if ext == "txt":
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return raw.decode(enc)
            except (UnicodeDecodeError, ValueError):
                continue
        return raw.decode("utf-8", errors="replace")

    elif ext == "pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages:
                    content = page.extract_text(x_tolerance=2, y_tolerance=3)
                    if content:
                        pages.append(content.strip())
            return "\n\n".join(pages) if pages else ""
        except ImportError:
            return "[PDF: instale pdfplumber — pip install pdfplumber]"
        except Exception as e:
            return f"[Erro na extração PDF: {e}]"

    elif ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n\n".join(parts)
        except ImportError:
            return "[DOCX: instale python-docx — pip install python-docx]"
        except Exception as e:
            return f"[Erro na extração DOCX: {e}]"

    return ""


def file_size_str(uploaded_file) -> str:
    """Tamanho legível do arquivo."""
    uploaded_file.seek(0, 2)
    size = uploaded_file.tell()
    uploaded_file.seek(0)
    if size < 1024:
        return f"{size} B"
    elif size < 1024 ** 2:
        return f"{size/1024:.1f} KB"
    return f"{size/1024**2:.1f} MB"


# ENGINE DE COMPARAÇÃO TEXTUAL

@st.cache_resource(show_spinner=False)
def _get_analyzer():
    """Instancia o motor de similaridade uma única vez por sessão do Streamlit."""
    from src.motor.similaridade import TextSimilarityAnalyzer
    return TextSimilarityAnalyzer()


def _find_common_snippets(text_a: str, text_b: str, n_words: int = 8) -> List[str]:
    """
    Encontra sequências de n palavras compartilhadas entre dois textos.
    Retorna as mais longas (até 5).
    """
    def ngrams(text: str, n: int):
        words = text.lower().split()
        return [" ".join(words[i:i+n]) for i in range(len(words) - n + 1)]

    shared = set(ngrams(text_a, n_words)) & set(ngrams(text_b, n_words))
    filtered = [s for s in shared if len(s) > 20]
    return sorted(filtered, key=len, reverse=True)[:5]


@dataclass
class PairResult:
    """Resultado de comparação de um par de documentos."""
    name_a: str
    name_b: str
    cosine_score: float
    jaccard_score: float
    embedding_score: float
    ensemble_score: float
    classification: str
    risk_level: str
    common_snippets: List[str] = field(default_factory=list)

    @property
    def ensemble_pct(self) -> float:
        return self.ensemble_score * 100


_RISK_LEVEL_BY_LABEL = {
    "ALTO RISCO": "high",
    "SIMILARIDADE ELEVADA": "elevado",
    "SIMILARIDADE MODERADA": "mod",
    "BAIXA SIMILARIDADE": "low",
}


def run_comparison(
    names: List[str],
    texts: List[str],
    progress_cb=None,
) -> List[PairResult]:
    """
    Executa comparação N×N entre todos os documentos usando o motor real
    (TextSimilarityAnalyzer). Cada par é processado individualmente pelo
    pipeline duplo (vetorização/embedding/jaccard) e classificado pelos
    limiares centralizados em config/parametros.py.

    Args:
        names: Nomes dos arquivos.
        texts: Textos extraídos.
        progress_cb: Callable(float) para atualizar barra de progresso.

    Returns:
        Lista de PairResult ordenada por ensemble_score decrescente.
    """
    analyzer = _get_analyzer()
    pairs = list(itertools.combinations(range(len(texts)), 2))
    results: List[PairResult] = []

    for step_idx, (i, j) in enumerate(pairs):
        if progress_cb:
            pct = 0.05 + 0.90 * (step_idx + 1) / max(len(pairs), 1)
            progress_cb(pct, f"Comparando: {names[i]} × {names[j]}")

        resultado = analyzer.analyze(texts[i], texts[j])

        if not resultado.is_success:
            logger.warning(
                "Falha ao comparar '%s' x '%s': %s",
                names[i], names[j], getattr(resultado, "message", "erro desconhecido"),
            )
            continue

        snippets = _find_common_snippets(texts[i], texts[j])

        results.append(PairResult(
            name_a=names[i],
            name_b=names[j],
            cosine_score=resultado.cosine_score,
            jaccard_score=resultado.jaccard_char_score,
            embedding_score=resultado.embedding_score,
            ensemble_score=resultado.ensemble_score,
            classification=resultado.classification,
            risk_level=_RISK_LEVEL_BY_LABEL.get(resultado.classification, "low"),
            common_snippets=snippets,
        ))

    if progress_cb:
        progress_cb(1.0, "Concluído.")

    return sorted(results, key=lambda r: r.ensemble_score, reverse=True)


# HELPERS DE RENDERIZAÇÃO HTML

def _score_bar_html(score: float) -> str:
    """Barra de progresso colorida inline."""
    pct = score * 100
    if score >= 0.80:
        color = "#F85149"
    elif score >= 0.60:
        color = "#D29922"
    elif score >= 0.35:
        color = "#00D4FF"
    else:
        color = "#3FB950"
    return f"""
    <div class="score-bar-wrap">
      <div class="score-bar-track">
        <div class="score-bar-fill" style="width:{pct:.1f}%;background:{color};"></div>
      </div>
      <div class="score-pct">{pct:.1f}%</div>
    </div>"""


def _risk_pill_html(label: str, risk: str) -> str:
    css = {"high": "risk-high", "elevado": "risk-elevado", "mod": "risk-mod", "low": "risk-low"}
    icons = {"high": "🔴", "elevado": "🟡", "mod": "🔵", "low": "🟢"}
    return (
        f'<span class="risk-pill {css.get(risk, "risk-low")}">'
        f'{icons.get(risk, "⚪")} {label}</span>'
    )



def _file_badge_html(ext: str) -> str:
    css = {"pdf": "badge-pdf", "docx": "badge-docx", "txt": "badge-txt"}
    return f'<span class="file-badge {css.get(ext, "badge-txt")}">{ext.upper()}</span>'


def _file_icon(ext: str) -> str:
    return {"pdf": "📄", "docx": "📝", "txt": "📃"}.get(ext, "📎")


# GRÁFICOS (ploty)

PLOTLY_THEME = dict(
    paper_bgcolor="#0D1117",
    plot_bgcolor="#161B22",
    font=dict(family="IBM Plex Mono, monospace", color="#8B949E", size=11),
    margin=dict(l=60, r=20, t=40, b=60),
)


def _build_heatmap(names: List[str], results: List[PairResult]) -> go.Figure:
    """Heatmap N×N de similaridade entre todos os documentos."""
    n = len(names)
    matrix = np.zeros((n, n))
    np.fill_diagonal(matrix, 1.0)

    name_idx = {name: i for i, name in enumerate(names)}
    for r in results:
        i, j = name_idx[r.name_a], name_idx[r.name_b]
        matrix[i, j] = r.ensemble_score
        matrix[j, i] = r.ensemble_score

    short = [n[:18] + "…" if len(n) > 18 else n for n in names]

    fig = go.Figure(go.Heatmap(
        z=matrix,
        x=short,
        y=short,
        colorscale=[
            [0.00, "rgb(13, 17, 23)"],
            [0.35, "rgba(14, 165, 233, 0.12)"],
            [0.60, "rgba(124, 58, 237, 0.50)"],
            [0.80, "rgb(248, 81, 73)"],
            [1.00, "rgb(255, 107, 107)"],
        ],
        zmin=0, zmax=1,
        text=np.round(matrix * 100, 1),
        texttemplate="%{text}%",
        textfont=dict(size=10, family="IBM Plex Mono"),
        hoverongaps=False,
        colorbar=dict(
            title=dict(
                text="Score",
                font=dict(color="#8B949E", family="IBM Plex Mono"),
            ),
            tickfont=dict(color="#8B949E", family="IBM Plex Mono"),
            bgcolor="#161B22",
            bordercolor="#30363D",
        ),
    ))

    fig.update_layout(
        title=dict(text="Matriz de Similaridade (Ensemble)", font=dict(color="#E6EDF3", size=13)),
        **PLOTLY_THEME,
        height=max(350, 80 * n),
        xaxis=dict(tickfont=dict(size=10, color="#8B949E"), gridcolor="#30363D"),
        yaxis=dict(tickfont=dict(size=10, color="#8B949E"), gridcolor="#30363D"),
    )
    return fig


def _build_radar(result: PairResult) -> go.Figure:
    """Radar chart com os 3 scores de um par."""
    categories = ["TF-IDF Cosseno", "Jaccard N-Gram", "Ensemble Híbrido"]
    values = [result.cosine_score, result.jaccard_score, result.ensemble_score]
    values_pct = [v * 100 for v in values]

    fig = go.Figure(go.Scatterpolar(
        r=values_pct + [values_pct[0]],
        theta=categories + [categories[0]],
        fill="toself",
        fillcolor="rgba(0, 212, 255, 0.12)",
        line=dict(color="#00D4FF", width=2),
        marker=dict(size=6, color="#00D4FF"),
    ))
    radar_theme = {k: v for k, v in PLOTLY_THEME.items() if k != "margin"}
    fig.update_layout(
        polar=dict(
            bgcolor="#161B22",
            radialaxis=dict(
                visible=True, range=[0, 100],
                tickfont=dict(size=9, color="#8B949E", family="IBM Plex Mono"),
                gridcolor="#30363D",
                linecolor="#30363D",
            ),
            angularaxis=dict(
                tickfont=dict(size=9, color="#E6EDF3", family="IBM Plex Mono"),
                linecolor="#30363D",
            ),
        ),
        **radar_theme,
        height=280,
        margin=dict(l=50, r=50, t=30, b=30),
    )
    return fig


def _build_bar_ranking(results: List[PairResult]) -> go.Figure:
    """Bar chart horizontal do ranking de pares."""
    labels = [f"{r.name_a[:12]}… × {r.name_b[:12]}…"
              if len(r.name_a) + len(r.name_b) > 28
              else f"{r.name_a} × {r.name_b}"
              for r in results]
    scores = [r.ensemble_pct for r in results]
    colors = []
    for r in results:
        if r.risk_level == "high":
            colors.append("#F85149")
        elif r.risk_level == "elevado":
            colors.append("#D29922")
        elif r.risk_level == "mod":
            colors.append("#00D4FF")
        else:
            colors.append("#3FB950")

    fig = go.Figure(go.Bar(
        x=scores,
        y=labels,
        orientation="h",
        marker_color=colors,
        text=[f"{s:.1f}%" for s in scores],
        textposition="outside",
        textfont=dict(size=10, family="IBM Plex Mono", color="#E6EDF3"),
    ))
    fig.update_layout(
        title=dict(text="Ranking de Similaridade por Par", font=dict(color="#E6EDF3", size=13)),
        xaxis=dict(
            title="Score (%)", range=[0, 115],
            tickfont=dict(size=9, color="#8B949E", family="IBM Plex Mono"),
            gridcolor="#30363D",
        ),
        yaxis=dict(
            tickfont=dict(size=10, color="#E6EDF3", family="IBM Plex Mono"),
            autorange="reversed",
        ),
        **PLOTLY_THEME,
        height=max(300, 60 * len(results)),
    )
    return fig


# EXPORTAÇÃO

def _export_csv(results: List[PairResult]) -> bytes:
    """Exporta resultados em CSV."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "Documento A", "Documento B",
        "Score TF-IDF Cosseno", "Score Jaccard",
        "Score Ensemble", "Classificação",
    ])
    for r in results:
        writer.writerow([
            r.name_a, r.name_b,
            f"{r.cosine_score:.4f}", f"{r.jaccard_score:.4f}",
            f"{r.ensemble_score:.4f}", r.classification,
        ])
    return buf.getvalue().encode("utf-8")


def _export_txt_report(results: List[PairResult], file_names: List[str]) -> bytes:
    """Exporta relatório técnico em TXT."""
    lines = [
        "=" * 70,
        "  SISTEMA DE AUDITORIA TEXTUAL — RELATÓRIO TÉCNICO",
        f"  Documentos analisados: {len(file_names)}",
        f"  Pares comparados: {len(results)}",
        "=" * 70,
        "",
    ]
    for i, r in enumerate(results, 1):
        lines += [
            f"[{i:02d}] {r.name_a}  ×  {r.name_b}",
            f"     Ensemble : {r.ensemble_pct:.2f}%  |  Cosseno: {r.cosine_score*100:.2f}%  |  Jaccard: {r.jaccard_score*100:.2f}%",
            f"     Status   : {r.classification}",
            "",
        ]
    return "\n".join(lines).encode("utf-8")


# SIDEBAR

def render_sidebar():
    """Painel lateral com configurações e instruções."""
    with st.sidebar:
        st.markdown("""
        <div style="font-family:'IBM Plex Mono';font-size:1rem;font-weight:700;
                    color:#00D4FF;margin-bottom:1.2rem;letter-spacing:-0.02em;">
          ⚙ Configuração
        </div>
        """, unsafe_allow_html=True)

        st.markdown(
            '<p style="font-size:0.72rem;color:#8B949E;font-family:\'IBM Plex Mono\';">'
            'LIMIARES DE CLASSIFICAÇÃO</p>',
            unsafe_allow_html=True
        )
        t_high = st.slider("Alto Risco (≥)", 0.50, 0.99, 0.80, 0.01,
                           help="Score acima deste limiar = ALTO RISCO")
        t_elev = st.slider("Similaridade Elevada (≥)", 0.30, 0.79, 0.60, 0.01)
        t_mod  = st.slider("Moderada (≥)", 0.10, 0.59, 0.35, 0.01)

        st.markdown("---")

        st.markdown(
            '<p style="font-size:0.72rem;color:#8B949E;font-family:\'IBM Plex Mono\';">'
            'OPÇÕES DE ANÁLISE</p>',
            unsafe_allow_html=True
        )
        show_snippets = st.checkbox("Exibir trechos comuns", value=True)
        show_radar    = st.checkbox("Gráfico radar por par", value=True)
        show_heatmap  = st.checkbox("Heatmap de similaridade", value=True)
        snippet_len   = st.slider("Tamanho do snippet (palavras)", 5, 15, 8)

        st.markdown("---")

        st.markdown("""
        <div class="tooltip-box">
          <b style="color:#E6EDF3;">Engine ativo</b><br>
          TF-IDF + Jaccard N-Gram<br>
          <span style="color:#00D4FF;">+ Embeddings</span> se sentence-transformers disponível
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("""
        <div style="font-size:0.70rem;color:#484F58;font-family:'IBM Plex Mono';line-height:1.8;">
          Vinicius de Oliveira Lima<br>
          github.com/68vinicius/TCC
        </div>
        """, unsafe_allow_html=True)

    return {
        "t_high": t_high,
        "t_elev": t_elev,
        "t_mod": t_mod,
        "show_snippets": show_snippets,
        "show_radar": show_radar,
        "show_heatmap": show_heatmap,
        "snippet_len": snippet_len,
    }


# MAIN APP

def main():
    # Injeta CSS global
    st.markdown(GLOBAL_CSS, unsafe_allow_html=True)

    # Sidebar
    cfg = render_sidebar()

    # HERO 
    st.markdown("""
    <div class="hero-container">
      <div class="hero-title">AUDIT<span>TEXT</span></div>
      <div class="hero-subtitle">
        Sistema Computacional de Auditoria Textual · Análise de Similaridade Documental
      </div>
      <div class="hero-badge">TF-IDF · Jaccard · Sentence Embeddings · Ensemble Híbrido</div>
    </div>
    """, unsafe_allow_html=True)

    # UPLOAD 
    st.markdown('<div class="section-label">01 · Upload de Documentos</div>',
                unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Arraste ou selecione 2 ou mais arquivos para comparação simultânea",
        type=["txt", "pdf", "docx"],
        accept_multiple_files=True,
        help="Formatos suportados: .txt, .pdf, .docx — mínimo 2 arquivos",
        label_visibility="visible",
    )

    # LISTA DE ARQUIVOS 
    if uploaded_files:
        st.markdown(
            f'<div class="section-label">02 · Arquivos Carregados '
            f'<span style="color:#8B949E;font-size:0.65rem;">({len(uploaded_files)} arquivo(s))</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

        col_files, col_summary = st.columns([3, 1])

        with col_files:
            for f in uploaded_files:
                ext = f.name.rsplit(".", 1)[-1].lower()
                size_str = file_size_str(f)
                icon = _file_icon(ext)
                badge = _file_badge_html(ext)
                # Hash curto para identificação
                f.seek(0)
                short_hash = hashlib.md5(f.read(512)).hexdigest()[:6].upper()
                f.seek(0)

                st.markdown(f"""
                <div class="file-card">
                  <span class="file-icon">{icon}</span>
                  <div class="file-info">
                    <div class="file-name">{f.name}</div>
                    <div class="file-meta">{size_str} &nbsp;·&nbsp; ID: {short_hash}</div>
                  </div>
                  {badge}
                </div>
                """, unsafe_allow_html=True)

        with col_summary:
            n_pdf  = sum(1 for f in uploaded_files if f.name.endswith(".pdf"))
            n_docx = sum(1 for f in uploaded_files if f.name.endswith(".docx"))
            n_txt  = sum(1 for f in uploaded_files if f.name.endswith(".txt"))
            n_pairs = len(uploaded_files) * (len(uploaded_files) - 1) // 2

            st.markdown(f"""
            <div class="metric-card cyan" style="margin-bottom:0.6rem;">
              <div class="metric-label">Total</div>
              <div class="metric-value">{len(uploaded_files)}</div>
              <div class="metric-sub">arquivos</div>
            </div>
            <div class="metric-card purple" style="margin-bottom:0.6rem;">
              <div class="metric-label">Pares</div>
              <div class="metric-value">{n_pairs}</div>
              <div class="metric-sub">comparações</div>
            </div>
            <div class="metric-card green">
              <div class="metric-label">Tipos</div>
              <div class="metric-value" style="font-size:0.9rem;line-height:1.6;">
                {"📄 " + str(n_pdf) + " PDF · " if n_pdf else ""}{"📝 " + str(n_docx) + " DOCX · " if n_docx else ""}{"📃 " + str(n_txt) + " TXT" if n_txt else ""}
              </div>
            </div>
            """, unsafe_allow_html=True)

    # VALIDAÇÃO E BOTÃO 
    st.markdown("---")

    if not uploaded_files:
        st.markdown("""
        <div style="text-align:center;padding:3rem 0;color:#484F58;font-family:'IBM Plex Mono';font-size:0.88rem;">
          ↑ Carregue pelo menos 2 arquivos para iniciar a auditoria
        </div>
        """, unsafe_allow_html=True)
        return

    if len(uploaded_files) < 2:
        st.warning("⚠ Carregue **pelo menos 2 arquivos** para comparação.")
        return

    col_btn, col_info = st.columns([2, 3])
    with col_btn:
        run = st.button(
            f"▶  INICIAR AUDITORIA  ({len(uploaded_files)} docs · {len(uploaded_files)*(len(uploaded_files)-1)//2} pares)",
            use_container_width=True,
        )

    with col_info:
        st.markdown("""
        <div class="tooltip-box" style="margin-top:0;">
          O sistema comparará <b style="color:#E6EDF3;">todos os pares possíveis</b>
          automaticamente. O score final é um ensemble ponderado de
          TF-IDF Cosseno, Jaccard N-Gram e Embeddings Semânticos.
        </div>
        """, unsafe_allow_html=True)

    if not run:
        return

    # PROCESSAMENTO 
    st.markdown('<div class="section-label">03 · Processando</div>', unsafe_allow_html=True)

    progress_bar = st.progress(0)
    status_text  = st.empty()

    def update_progress(pct: float, msg: str):
        progress_bar.progress(pct)
        status_text.markdown(
            f'<span style="font-family:\'IBM Plex Mono\';font-size:0.78rem;color:#8B949E;">'
            f'⟳ {msg}</span>',
            unsafe_allow_html=True,
        )

    update_progress(0.02, "Extraindo textos dos documentos...")
    names, texts = [], []
    extraction_errors = []

    for f in uploaded_files:
        update_progress(0.02, f"Extraindo: {f.name}")
        try:
            text = extract_text(f)
            if len(text.strip()) < 20:
                extraction_errors.append(f.name)
                continue
            names.append(f.name)
            texts.append(text)
        except Exception as e:
            extraction_errors.append(f"{f.name} ({e})")

    if extraction_errors:
        st.warning(f"⚠ Não foi possível extrair texto de: {', '.join(extraction_errors)}")

    if len(texts) < 2:
        st.error("Erro: menos de 2 documentos com texto extraível.")
        progress_bar.empty()
        status_text.empty()
        return

    # Comparação
    t_start = time.time()
    results = run_comparison(names, texts, progress_cb=update_progress)
    elapsed = time.time() - t_start

    progress_bar.empty()
    status_text.empty()

    # MÉTRICAS GLOBAIS 
    st.markdown('<div class="section-label">04 · Resumo Executivo</div>',
                unsafe_allow_html=True)

    avg_score = float(np.mean([r.ensemble_score for r in results]))
    max_score = max(r.ensemble_score for r in results)
    n_risk    = sum(1 for r in results if r.risk_level == "high")

    st.markdown(f"""
    <div class="metric-grid">
      <div class="metric-card cyan">
        <div class="metric-label">Score Médio</div>
        <div class="metric-value">{avg_score*100:.1f}%</div>
        <div class="metric-sub">ensemble híbrido</div>
      </div>
      <div class="metric-card {"red" if max_score >= 0.80 else "yellow" if max_score >= 0.60 else "cyan"}">
        <div class="metric-label">Score Máximo</div>
        <div class="metric-value">{max_score*100:.1f}%</div>
        <div class="metric-sub">par mais similar</div>
      </div>
      <div class="metric-card purple">
        <div class="metric-label">Pares Analisados</div>
        <div class="metric-value">{len(results)}</div>
        <div class="metric-sub">{len(names)} documentos</div>
      </div>
      <div class="metric-card {"red" if n_risk > 0 else "green"}">
        <div class="metric-label">Alto Risco</div>
        <div class="metric-value">{n_risk}</div>
        <div class="metric-sub">{"par(es) crítico(s)" if n_risk > 0 else "nenhum detectado"}</div>
      </div>
      <div class="metric-card yellow">
        <div class="metric-label">Latência</div>
        <div class="metric-value">{elapsed:.2f}s</div>
        <div class="metric-sub">tempo total</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # RANKING 
    st.markdown('<div class="section-label">05 · Ranking de Similaridade</div>',
                unsafe_allow_html=True)

    for rank, r in enumerate(results, 1):
        rank_cls = f"rank-{rank}" if rank <= 3 else ""
        bar_html = _score_bar_html(r.ensemble_score)
        pill_html = _risk_pill_html(r.classification, r.risk_level)

        st.markdown(f"""
        <div class="result-row">
          <div class="rank-badge {rank_cls}">#{rank:02d}</div>
          <div class="pair-names">
            {r.name_a}<br>
            <span>×&nbsp;&nbsp;</span>{r.name_b}
          </div>
          {bar_html}
          {pill_html}
        </div>
        """, unsafe_allow_html=True)

    # TABELA DETALHADA 
    st.markdown('<div class="section-label">06 · Tabela Analítica Detalhada</div>',
                unsafe_allow_html=True)

    df_data = {
        "Rank":          [f"#{i+1:02d}" for i in range(len(results))],
        "Documento A":   [r.name_a for r in results],
        "Documento B":   [r.name_b for r in results],
        "TF-IDF Cos %":  [f"{r.cosine_score*100:.2f}" for r in results],
        "Jaccard %":     [f"{r.jaccard_score*100:.2f}" for r in results],
        "Ensemble %":    [f"{r.ensemble_pct:.2f}" for r in results],
        "Classificação": [r.classification for r in results],
    }
    df = pd.DataFrame(df_data)
    st.dataframe(df, use_container_width=True, hide_index=True)

    # GRÁFICOS 
    st.markdown('<div class="section-label">07 · Visualizações</div>',
                unsafe_allow_html=True)

    if cfg["show_heatmap"] and len(names) >= 2:
        fig_heat = _build_heatmap(names, results)
        st.plotly_chart(fig_heat, use_container_width=True)

    if len(results) > 0:
        fig_bar = _build_bar_ranking(results)
        st.plotly_chart(fig_bar, use_container_width=True)

    # ANÁLISE POR PAR (EXPANDERS) 
    st.markdown('<div class="section-label">08 · Análise Individual por Par</div>',
                unsafe_allow_html=True)

    for rank, r in enumerate(results, 1):
        label_short = f"#{rank:02d} · {r.name_a} × {r.name_b} — {r.ensemble_pct:.1f}%  {r.classification}"

        with st.expander(label_short, expanded=(rank == 1)):

            left, right = st.columns([1, 1])

            with left:
                st.markdown(f"""
                <div style="margin-bottom:0.8rem;">
                  <div class="metric-label">SCORES DETALHADOS</div>
                </div>
                <table style="width:100%;font-family:'IBM Plex Mono';font-size:0.80rem;
                              border-collapse:collapse;">
                  <tr style="border-bottom:1px solid #30363D;">
                    <td style="padding:0.4rem 0;color:#8B949E;">TF-IDF Cosseno</td>
                    <td style="text-align:right;color:#E6EDF3;font-weight:600;">
                      {r.cosine_score*100:.2f}%
                    </td>
                  </tr>
                  <tr style="border-bottom:1px solid #30363D;">
                    <td style="padding:0.4rem 0;color:#8B949E;">Jaccard N-Gram</td>
                    <td style="text-align:right;color:#E6EDF3;font-weight:600;">
                      {r.jaccard_score*100:.2f}%
                    </td>
                  </tr>
                  <tr>
                    <td style="padding:0.4rem 0;color:#00D4FF;font-weight:600;">Ensemble Híbrido</td>
                    <td style="text-align:right;color:#00D4FF;font-weight:700;">
                      {r.ensemble_pct:.2f}%
                    </td>
                  </tr>
                </table>
                <div style="margin-top:1rem;">
                  {_risk_pill_html(r.classification, r.risk_level)}
                </div>
                """, unsafe_allow_html=True)

            with right:
                if cfg["show_radar"]:
                    fig_radar = _build_radar(r)
                    st.plotly_chart(fig_radar, use_container_width=True)

            if cfg["show_snippets"]:
                snippets = _find_common_snippets(
                    texts[names.index(r.name_a)],
                    texts[names.index(r.name_b)],
                    n_words=cfg["snippet_len"],
                )
                if snippets:
                    st.markdown(
                        f'<div class="metric-label" style="margin-top:0.8rem;">'
                        f'TRECHOS COMUNS DETECTADOS ({len(snippets)})</div>',
                        unsafe_allow_html=True,
                    )
                    for snip in snippets:
                        st.markdown(
                            f'<div class="snippet-box">…{snip}…</div>',
                            unsafe_allow_html=True,
                        )
                else:
                    st.markdown(
                        '<div class="snippet-box" style="color:#484F58;">'
                        'Nenhuma sequência textual compartilhada detectada com o comprimento configurado.'
                        '</div>',
                        unsafe_allow_html=True,
                    )

    # EXPORTAÇÃO 
    st.markdown('<div class="section-label">09 · Exportar Resultados</div>',
                unsafe_allow_html=True)

    col_csv, col_txt, col_space = st.columns([1, 1, 2])

    with col_csv:
        st.download_button(
            label="⬇ Baixar CSV",
            data=_export_csv(results),
            file_name="auditoria_resultados.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with col_txt:
        st.download_button(
            label="⬇ Baixar Relatório TXT",
            data=_export_txt_report(results, names),
            file_name="auditoria_relatorio.txt",
            mime="text/plain",
            use_container_width=True,
        )

    # FOOTER 
    st.markdown("---")
    st.markdown(f"""
    <div style="text-align:center;font-family:'IBM Plex Mono';font-size:0.70rem;
                color:#484F58;padding:1rem 0 0.5rem;line-height:2;">
      Sistema Computacional de Auditoria Textual &nbsp;·&nbsp;
      TCC 2025 &nbsp;·&nbsp; Vinicius de Oliveira Lima<br>
      Engine: TF-IDF + Jaccard N-Gram + Sentence Embeddings &nbsp;·&nbsp;
      {len(results)} pares analisados em {elapsed:.2f}s
    </div>
    """, unsafe_allow_html=True)



if __name__ == "__main__":
    main()