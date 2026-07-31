from __future__ import annotations

import io
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, Union

import pdfplumber
from charset_normalizer import from_bytes
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

FileInput = Union[io.BytesIO, str, Path]


# DATA 

@dataclass
class ExtractedDocument:
    """
    Resultado estruturado da extração documental.

    Separa o conteúdo textual de metadados para permitir análises
    distintas sobre cada componente.
    """
    raw_text: str
    filename: str
    extension: str
    char_count: int = field(init=False)
    word_count: int = field(init=False)
    page_count: Optional[int] = None
    metadata: dict = field(default_factory=dict)
    extraction_warnings: list[str] = field(default_factory=list)

    def __post_init__(self):
        self.char_count = len(self.raw_text)
        self.word_count = len(self.raw_text.split())

    @property
    def is_empty(self) -> bool:
        return self.word_count < 10

    @property
    def is_too_short(self) -> bool:
        return self.char_count < 50


# ABSTRACT HANDLER

class DocumentHandler(ABC):
    """Interface do padrão Factory Method para extratores."""

    @abstractmethod
    def extract(self, file: FileInput, filename: str) -> ExtractedDocument:
        ...

    def _make_doc(
        self,
        text: str,
        filename: str,
        ext: str,
        **kwargs
    ) -> ExtractedDocument:
        return ExtractedDocument(
            raw_text=text,
            filename=filename,
            extension=ext,
            **kwargs
        )


# PDF HANDLER

class PDFHandler(DocumentHandler):
    """
    Extração de PDF via pdfplumber com fallback e heurística multi-coluna.

    pdfplumber supera pypdf para extração de texto com layout complexo
    porque usa heurísticas de agrupamento de caracteres por posição x,y,
    crítico para documentos acadêmicos com múltiplas colunas.
    """

    def extract(self, file: FileInput, filename: str) -> ExtractedDocument:
        warnings: list[str] = []
        pages_text: list[str] = []
        page_count = 0

        try:
            with pdfplumber.open(file) as pdf:
                page_count = len(pdf.pages)
                metadata = pdf.metadata or {}

                for i, page in enumerate(pdf.pages):
                    try:
                        content = page.extract_text(
                            x_tolerance=2,
                            y_tolerance=3,
                            layout=True,
                        )
                        if content and content.strip():
                            pages_text.append(content.strip())
                        else:
                            warnings.append(f"Página {i+1}: sem texto extraído (possivelmente imagem).")
                    except Exception as e:
                        warnings.append(f"Página {i+1}: erro de extração — {e}")
                        logger.warning("Erro na página %d do PDF '%s': %s", i + 1, filename, e)

        except Exception as e:
            logger.exception("Falha crítica ao abrir PDF '%s'", filename)
            raise ValueError(f"Não foi possível abrir o PDF '{filename}': {e}") from e

        full_text = "\n\n".join(pages_text)

        if not full_text.strip():
            warnings.append("Nenhum texto extraído. O PDF pode ser baseado em imagens (OCR necessário).")

        return self._make_doc(
            text=full_text,
            filename=filename,
            ext="pdf",
            page_count=page_count,
            metadata={k: str(v) for k, v in metadata.items() if v},
            extraction_warnings=warnings,
        )


# DOCX HANDLER

class DocxHandler(DocumentHandler):
    """
    Extração de DOCX preservando estrutura de parágrafos e tabelas.

    O sistema original ignorava tabelas, que em TCCs frequentemente
    contêm conteúdo relevante (dados experimentais, comparações).
    """

    def extract(self, file: FileInput, filename: str) -> ExtractedDocument:
        warnings: list[str] = []
        sections: list[str] = []

        try:
            doc = DocxDocument(file)

            # Parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text.strip())

            # Tabelas (conteúdo relevante ignorado pelo sistema original)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        sections.append(row_text)

            # Metadados do documento
            props = doc.core_properties
            metadata = {}
            for attr in ("author", "title", "subject", "created", "modified"):
                val = getattr(props, attr, None)
                if val:
                    metadata[attr] = str(val)

        except Exception as e:
            logger.exception("Falha ao processar DOCX '%s'", filename)
            raise ValueError(f"Não foi possível processar o DOCX '{filename}': {e}") from e

        full_text = "\n\n".join(sections)
        return self._make_doc(
            text=full_text,
            filename=filename,
            ext="docx",
            metadata=metadata,
            extraction_warnings=warnings,
        )


# TXT HANDLER

class TextHandler(DocumentHandler):
    """
    Extração de TXT com detecção automática de encoding.

    charset-normalizer detecta encoding com alta acurácia para
    português, crucial para documentos com caracteres especiais (ã, ç, etc.)
    exportados de sistemas legados.
    """

    def extract(self, file: FileInput, filename: str) -> ExtractedDocument:
        warnings: list[str] = []

        if isinstance(file, (str, Path)):
            with open(file, "rb") as f:
                raw = f.read()
        else:
            raw = file.read()

        result = from_bytes(raw)
        best = result.best()

        if best is None:
            raise ValueError(f"Não foi possível detectar encoding do arquivo '{filename}'.")

        encoding = best.encoding
        confidence = best.encoding_confidence if hasattr(best, "encoding_confidence") else "?"

        if str(confidence) != "?" and float(str(confidence)) < 0.7:
            warnings.append(
                f"Confiança baixa na detecção de encoding ({confidence:.0%}). "
                "Verifique se o arquivo está corrompido."
            )

        text = str(best)
        return self._make_doc(
            text=text,
            filename=filename,
            ext="txt",
            metadata={"detected_encoding": encoding},
            extraction_warnings=warnings,
        )


# FACTORY

class DocumentFactory:
    """
    Factory Method para instanciação de handlers por extensão.

    CORREÇÃO DO SISTEMA ORIGINAL:
    O original usava instâncias singleton compartilhadas no dict _map,
    o que pode causar race conditions em uso concorrente (ex: múltiplos
    usuários simultâneos no Streamlit). Esta versão instancia handlers
    sob demanda (stateless por design).
    """

    _REGISTRY: dict[str, type[DocumentHandler]] = {
        "pdf": PDFHandler,
        "docx": DocxHandler,
        "txt": TextHandler,
    }

    @classmethod
    def extract(cls, file: FileInput, filename: str) -> ExtractedDocument:
        """
        Extrai texto e metadados de um arquivo.

        Args:
            file: File-like object (Streamlit UploadedFile, BytesIO, Path).
            filename: Nome do arquivo com extensão.

        Returns:
            ExtractedDocument com texto e metadados.

        Raises:
            ValueError: Extensão não suportada ou falha de extração.
        """
        ext = Path(filename).suffix.lstrip(".").lower()
        handler_cls = cls._REGISTRY.get(ext)

        if handler_cls is None:
            supported = ", ".join(f".{e}" for e in cls._REGISTRY)
            raise ValueError(
                f"Extensão '.{ext}' não suportada. Formatos aceitos: {supported}"
            )

        handler = handler_cls()
        doc = handler.extract(file, filename)

        if doc.is_empty:
            raise ValueError(
                f"O documento '{filename}' não contém texto suficiente para análise "
                f"(mínimo: 10 palavras, encontradas: {doc.word_count})."
            )

        if doc.extraction_warnings:
            for w in doc.extraction_warnings:
                logger.warning("[%s] %s", filename, w)

        logger.info(
            "Documento '%s' extraído: %d chars, %d palavras, %d páginas.",
            filename,
            doc.char_count,
            doc.word_count,
            doc.page_count or 0,
        )

        return doc

    @classmethod
    def supported_extensions(cls) -> list[str]:
        return list(cls._REGISTRY.keys())